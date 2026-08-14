# -*- coding: utf-8 -*-
import streamlit as st
import json
import io
import re
import time
from pathlib import Path
from datetime import datetime

from mufredat_verisi import MUFREDAT
from mufredat_baglam import UNITE_BAGLAM, CIKTI_BAGLAM

try:
    import anthropic
    ANTHROPIC_MEVCUT = True
except ImportError:
    ANTHROPIC_MEVCUT = False

try:
    import google.generativeai as genai
    GEMINI_MEVCUT = True
except ImportError:
    GEMINI_MEVCUT = False

try:
    from pypdf import PdfReader
    PYPDF_MEVCUT = True
except ImportError:
    PYPDF_MEVCUT = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_MEVCUT = True
except ImportError:
    DOCX_MEVCUT = False

try:
    from openai import OpenAI  # DeepSeek OpenAI-uyumlu API kullanır
    OPENAI_MEVCUT = True
except ImportError:
    OPENAI_MEVCUT = False

# ==========================================
# SABİTLER VE KLASÖRLER
# ==========================================
MAKS_KAYNAK_KARAKTER_VARSAYILAN = 150000
UYGULAMA_KLASORU = Path(__file__).parent
KUTUPHANE_KLASORU = UYGULAMA_KLASORU / "kaynak_kutuphane"
SORU_HAVUZU_KLASORU = UYGULAMA_KLASORU / "soru_havuzu"
KUTUPHANE_KLASORU.mkdir(exist_ok=True)
SORU_HAVUZU_KLASORU.mkdir(exist_ok=True)
SORU_HAVUZU_DOSYA = SORU_HAVUZU_KLASORU / "havuz.json"

ZORLUK_SECENEKLERI = ["Kolay", "Orta", "Zor"]

# ------------------------------------------
# MODEL ADLARI — TEK YERDEN YÖNETİLİR
# Sağlayıcılar zaman zaman modelleri kullanımdan kaldırır (deprecate).
# Bir model kapatılırsa 404 alırsınız; sadece aşağıdaki listeyi güncelleyin.
# Güncel adlar için:
#   Gemini: https://ai.google.dev/gemini-api/docs/models
#   Claude: https://docs.claude.com
# NOT: gemini-2.0-flash 1 Haziran 2026'da kapatıldı; yerine gemini-3.5-flash.
# ------------------------------------------
GEMINI_MODELLERI = ["gemini-3.5-flash", "gemini-2.5-pro", "gemini-2.5-flash"]
CLAUDE_MODELLERI = ["claude-sonnet-5", "claude-opus-4-8", "claude-haiku-4-5-20251001"]
# DeepSeek OpenAI-uyumlu API kullanır. Katı dakikalık istek limiti yoktur, çok ucuzdur.
# Model adları için: https://api-docs.deepseek.com  (eski deepseek-chat/reasoner emekli oldu)
DEEPSEEK_MODELLERI = ["deepseek-v4-flash", "deepseek-v4-pro"]
DEEPSEEK_BASE_URL = "https://api.deepseek.com"

# ==========================================
# SAYFA AYARLARI
# ==========================================
st.set_page_config(
    page_title="TYMM 11. Sınıf Tarih Soru Üreteci",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# YARDIMCI — KAYNAK KÜTÜPHANESİ
# ==========================================

def pdf_metin_cikar(dosya_yolu):
    reader = PdfReader(dosya_yolu)
    parcalar = []
    for sayfa in reader.pages:
        try:
            parcalar.append(sayfa.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parcalar).strip()


def txt_metin_cikar(dosya_yolu):
    ham = Path(dosya_yolu).read_bytes()
    for kodlama in ("utf-8", "windows-1254", "iso-8859-9", "latin-1"):
        try:
            return ham.decode(kodlama).strip()
        except (UnicodeDecodeError, AttributeError):
            continue
    return ham.decode("utf-8", errors="ignore").strip()


@st.cache_data(show_spinner=False)
def dosyadan_metin_cikar(dosya_yolu_str, degisiklik_zamani):
    dosya_yolu = Path(dosya_yolu_str)
    if dosya_yolu.suffix.lower() == ".pdf":
        if not PYPDF_MEVCUT:
            return ""
        return pdf_metin_cikar(dosya_yolu)
    return txt_metin_cikar(dosya_yolu)


def dosyayi_kutuphaneye_kaydet(yuklenen_dosya):
    hedef = KUTUPHANE_KLASORU / yuklenen_dosya.name
    sayac = 1
    while hedef.exists():
        hedef = KUTUPHANE_KLASORU / f"{Path(yuklenen_dosya.name).stem}_{sayac}{Path(yuklenen_dosya.name).suffix}"
        sayac += 1
    hedef.write_bytes(yuklenen_dosya.getbuffer())
    return hedef


def kutuphaneyi_diskten_yukle():
    tum_metin, dosya_adlari = [], []
    for dosya_yolu in sorted(KUTUPHANE_KLASORU.glob("*")):
        if dosya_yolu.suffix.lower() not in (".pdf", ".txt"):
            continue
        try:
            metin = dosyadan_metin_cikar(str(dosya_yolu), dosya_yolu.stat().st_mtime)
        except Exception as e:
            st.error(f"'{dosya_yolu.name}' okunurken hata: {e}")
            continue
        if metin:
            tum_metin.append(f"### Kaynak: {dosya_yolu.name}\n{metin}")
            dosya_adlari.append(dosya_yolu.name)
        else:
            st.warning(f"'{dosya_yolu.name}' içinden metin çıkarılamadı (taranmış görsel PDF olabilir).")
    return "\n\n---\n\n".join(tum_metin), dosya_adlari


def yuklenen_dosyalari_isle(dosyalar):
    for dosya in dosyalar:
        try:
            dosyayi_kutuphaneye_kaydet(dosya)
        except Exception as e:
            st.error(f"'{dosya.name}' diske kaydedilirken hata: {e}")


# ==========================================
# YARDIMCI — SORU HAVUZU
# ==========================================

def havuzu_yukle():
    if SORU_HAVUZU_DOSYA.exists():
        try:
            return json.loads(SORU_HAVUZU_DOSYA.read_text(encoding="utf-8"))
        except Exception:
            return []
    return []


def havuza_kaydet(kayit):
    havuz = havuzu_yukle()
    havuz.insert(0, kayit)
    SORU_HAVUZU_DOSYA.write_text(json.dumps(havuz, ensure_ascii=False, indent=2), encoding="utf-8")


def havuzdan_sil(kayit_id):
    havuz = [k for k in havuzu_yukle() if k.get("id") != kayit_id]
    SORU_HAVUZU_DOSYA.write_text(json.dumps(havuz, ensure_ascii=False, indent=2), encoding="utf-8")


# ==========================================
# YARDIMCI — WORD (.docx) DIŞA AKTARMA
# ==========================================

def _markdown_run_ekle(paragraf, satir):
    """Bir satırdaki **kalın** parçaları Word run'larına çevirir."""
    for parca in re.split(r'(\*\*.+?\*\*)', satir):
        if not parca:
            continue
        if parca.startswith("**") and parca.endswith("**") and len(parca) > 4:
            paragraf.add_run(parca[2:-2]).bold = True
        else:
            paragraf.add_run(parca)


def _markdown_word_ekle(dokuman, markdown_metin):
    """Markdown içeriği (soru metni) Word paragraflarına dönüştürür.
       Başlık (#), madde/numaralı liste, kalın (**) ve yatay çizgi desteklenir."""
    for ham in (markdown_metin or "").split("\n"):
        satir = ham.rstrip()
        if not satir.strip():
            continue
        if satir.strip() in ("---", "***", "___"):
            continue
        baslik_esle = re.match(r'^(#{1,6})\s+(.*)', satir)
        if baslik_esle:
            seviye = min(len(baslik_esle.group(1)), 4)
            dokuman.add_heading(baslik_esle.group(2).strip(), level=seviye)
            continue
        madde_esle = re.match(r'^\s*[-*+]\s+(.*)', satir)
        if madde_esle:
            p = dokuman.add_paragraph(style="List Bullet")
            _markdown_run_ekle(p, madde_esle.group(1))
            continue
        num_esle = re.match(r'^\s*(\d+)[\.\)]\s+(.*)', satir)
        if num_esle:
            p = dokuman.add_paragraph(style="List Number")
            _markdown_run_ekle(p, num_esle.group(2))
            continue
        p = dokuman.add_paragraph()
        _markdown_run_ekle(p, satir)


def _kayit_word_ekle(dokuman, kayit):
    """Tek bir havuz kaydını (meta + içerik) belgeye ekler."""
    parcalar = [kayit.get("cikti_kod", ""), kayit.get("soru_kategorisi", ""),
                kayit.get("zorluk", "")]
    if kayit.get("soru_sayisi"):
        parcalar.append(f"{kayit.get('soru_sayisi')} soru")
    bilgi = " · ".join(str(x) for x in parcalar if str(x).strip())
    dokuman.add_heading(bilgi or "Soru Kaydı", level=1)

    for etiket, anahtar in (("Ünite", "unite"), ("Öğrenme Çıktısı", "cikti_tam")):
        deger = kayit.get(anahtar, "")
        if deger:
            p = dokuman.add_paragraph()
            p.add_run(f"{etiket}: ").bold = True
            p.add_run(str(deger))
    if kayit.get("surecler"):
        p = dokuman.add_paragraph()
        p.add_run("Alt başlıklar: ").bold = True
        p.add_run("; ".join(kayit["surecler"]))
    alt = []
    if kayit.get("model"):
        alt.append(f"Model: {kayit['model']}")
    if kayit.get("zaman"):
        alt.append(f"Üretim: {kayit['zaman']}")
    if alt:
        p = dokuman.add_paragraph()
        p.add_run(" · ".join(alt)).italic = True

    dokuman.add_paragraph()
    _markdown_word_ekle(dokuman, kayit.get("icerik", ""))


def tek_kayit_word(kayit):
    """Tek kaydı Word belgesine çevirir; BytesIO döner."""
    dokuman = Document()
    dokuman.add_heading("TYMM 11. Sınıf Tarih — Soru", level=0)
    _kayit_word_ekle(dokuman, kayit)
    tampon = io.BytesIO()
    dokuman.save(tampon)
    tampon.seek(0)
    return tampon


def coklu_kayit_word(kayitlar, baslik="TYMM 11. Sınıf Tarih — Soru Havuzu"):
    """Birden çok kaydı tek Word belgesinde toplar; her kayıt yeni sayfada."""
    dokuman = Document()
    dokuman.add_heading(baslik, level=0)
    dokuman.add_paragraph(
        f"Toplam {len(kayitlar)} kayıt · Dışa aktarma: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    for i, kayit in enumerate(kayitlar):
        if i > 0:
            dokuman.add_page_break()
        _kayit_word_ekle(dokuman, kayit)
    tampon = io.BytesIO()
    dokuman.save(tampon)
    tampon.seek(0)
    return tampon


def ham_metin_word(markdown_metin, baslik="TYMM 11. Sınıf Tarih — Üretilen Sorular", ust_bilgi=None):
    """Serbest Markdown metni (henüz havuza kaydedilmemiş üretim) Word'e çevirir."""
    dokuman = Document()
    dokuman.add_heading(baslik, level=0)
    if ust_bilgi:
        p = dokuman.add_paragraph()
        p.add_run(ust_bilgi).italic = True
        dokuman.add_paragraph()
    _markdown_word_ekle(dokuman, markdown_metin)
    tampon = io.BytesIO()
    dokuman.save(tampon)
    tampon.seek(0)
    return tampon


# ==========================================
# YARDIMCI — AI ÇAĞRILARI
# ==========================================

def _bekleme_suresi_bul(mesaj, varsayilan=35):
    """429 hata metnindeki retry_delay saniyesini yakalar; yoksa varsayılanı döner."""
    eslesme = re.search(r'retry.{0,25}?(\d+)', mesaj, re.IGNORECASE | re.DOTALL)
    if eslesme:
        try:
            return int(eslesme.group(1)) + 2
        except ValueError:
            pass
    return varsayilan


def soru_uret_api(api_key, model, prompt, max_deneme=3, ilerleme=None):
    """Claude ile üretim. 429/aşırı yük (overloaded) durumunda otomatik yeniden dener."""
    client = anthropic.Anthropic(api_key=api_key)
    for deneme in range(max_deneme):
        try:
            response = client.messages.create(
                model=model, max_tokens=16000,
                messages=[{"role": "user", "content": prompt}]
            )
            metin = "\n".join(b.text for b in response.content if b.type == "text")
            if response.stop_reason == "max_tokens":
                metin += "\n\n---\n⚠️ **UYARI:** Yanıt token sınırına takılıp yarım kalmış olabilir. Soru sayısını azaltıp tekrar deneyin."
            return metin
        except Exception as e:
            mesaj = str(e)
            gecici = ("429" in mesaj or "rate_limit" in mesaj.lower()
                      or "overloaded" in mesaj.lower() or "529" in mesaj)
            if gecici and deneme < max_deneme - 1:
                bekleme = _bekleme_suresi_bul(mesaj)
                if ilerleme:
                    ilerleme(f"Kota/yoğunluk limiti — {bekleme} sn beklenip tekrar denenecek "
                             f"({deneme + 1}/{max_deneme - 1})...")
                time.sleep(bekleme)
                continue
            raise
    raise RuntimeError("Kota limiti nedeniyle üretim başarısız oldu. Lütfen biraz sonra tekrar deneyin.")


def soru_uret_gemini(api_key, model, prompt, max_deneme=3, ilerleme=None):
    """Gemini ile üretim. 429 kota durumunda retry_delay kadar bekleyip otomatik yeniden dener."""
    genai.configure(api_key=api_key)
    model_obj = genai.GenerativeModel(model)
    for deneme in range(max_deneme):
        try:
            response = model_obj.generate_content(
                prompt, generation_config=genai.types.GenerationConfig(max_output_tokens=16000)
            )
            metin = response.text
            try:
                if response.candidates[0].finish_reason == 2:
                    metin += "\n\n---\n⚠️ **UYARI:** Yanıt token sınırına takılıp yarım kalmış olabilir. Soru sayısını azaltıp tekrar deneyin."
            except (IndexError, AttributeError):
                pass
            return metin
        except Exception as e:
            mesaj = str(e)
            if ("429" in mesaj or "quota" in mesaj.lower() or "rate" in mesaj.lower()) and deneme < max_deneme - 1:
                bekleme = _bekleme_suresi_bul(mesaj)
                if ilerleme:
                    ilerleme(f"Gemini kota limiti — {bekleme} sn beklenip tekrar denenecek "
                             f"({deneme + 1}/{max_deneme - 1})...")
                time.sleep(bekleme)
                continue
            raise
    raise RuntimeError("Kota limiti nedeniyle üretim başarısız oldu. Lütfen biraz sonra tekrar deneyin.")


def soru_uret_deepseek(api_key, model, prompt, max_deneme=3, ilerleme=None):
    """DeepSeek ile üretim (OpenAI-uyumlu API). Yoğunluk/geçici hata durumunda yeniden dener."""
    client = OpenAI(api_key=api_key, base_url=DEEPSEEK_BASE_URL)
    for deneme in range(max_deneme):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=8000,
            )
            metin = response.choices[0].message.content or ""
            try:
                if response.choices[0].finish_reason == "length":
                    metin += "\n\n---\n⚠️ **UYARI:** Yanıt token sınırına takılıp yarım kalmış olabilir. Soru sayısını azaltıp tekrar deneyin."
            except (IndexError, AttributeError):
                pass
            return metin
        except Exception as e:
            mesaj = str(e)
            gecici = ("429" in mesaj or "rate" in mesaj.lower() or "503" in mesaj
                      or "overload" in mesaj.lower() or "timeout" in mesaj.lower())
            if gecici and deneme < max_deneme - 1:
                bekleme = _bekleme_suresi_bul(mesaj, varsayilan=20)
                if ilerleme:
                    ilerleme(f"DeepSeek yoğunluk/limit — {bekleme} sn beklenip tekrar denenecek "
                             f"({deneme + 1}/{max_deneme - 1})...")
                time.sleep(bekleme)
                continue
            raise
    raise RuntimeError("DeepSeek üretimi başarısız oldu. Lütfen biraz sonra tekrar deneyin.")

def build_prompt(unite, cikti_kod, cikti_tam, surec_metinleri, soru_kategorisi,
                 zorluk, soru_sayisi, kaynak_metin="", ek_baglam=""):
    baglam_temelli = (soru_kategorisi == "Bağlam Temelli")
    surec_listesi = "\n".join(f"   - {s}" for s in surec_metinleri)

    # --- Öğretim programı bağlamı (mufredat_baglam.py) ---
    ub = UNITE_BAGLAM.get(unite, {})
    cb = CIKTI_BAGLAM.get(cikti_kod, {})
    kavramlar = ", ".join(ub.get("anahtar_kavramlar", []))
    beceriler = "; ".join(ub.get("alan_becerileri", []))
    degerler = ", ".join(ub.get("degerler", []))
    olcumlenen_beceri = cb.get("olcumlenen_beceri", "")
    somut_icerik = cb.get("somut_icerik", [])
    somut_listesi = "\n".join(f"   - {s}" for s in somut_icerik)

    program_blok = f"""
### 📖 ÖĞRETİM PROGRAMI BAĞLAMI (TYMM — bu çıktıya özgü, ZORUNLU KULLANIM):
- **Bu çıktının ölçtüğü merkezî beceri:** {olcumlenen_beceri}
- **Ünitenin anahtar kavramları:** {kavramlar}
- **İlgili alan becerileri:** {beceriler}
- **İşlenen değerler:** {degerler}
- **Bu çıktı işlenirken programın öngördüğü SOMUT tarihsel içerik (olay, antlaşma, kişi, olgu):**
{somut_listesi}

**ZORUNLU:** Sorular ve bağlam metinleri, yukarıdaki SOMUT tarihsel içerikten beslenmelidir. Soyut/genel ifadeler yerine bu listedeki gerçek olay, antlaşma ve kişilere dayan. Birden fazla soru üretiyorsan, HER SORU bu somut içerikten FARKLI bir öğeyi merkeze alsın; hepsi aynı olayı/temayı tekrarlamasın. Böylece sorular hem müfredata birebir uyar hem de birbirinden gerçekten farklılaşır.
"""

    uslup_blok = """
0. **Yazım Üslubu (ZORUNLU):** Metinler deneyimli bir tarih öğretmeni tarafından kaleme alınmış gibi doğal, akıcı ve özgün bir Türkçeyle yazılmalı; yapay zekâ klişelerinden ("Günümüzde...", "Tarih boyunca...", "Bilindiği gibi...") ve şablon cümlelerden kaçınılmalıdır. Her bağlam metni kendine özgü olmalı, seri üretim hissi vermemelidir.
"""

    dogruluk_blok = """
0b. **Tarihsel Doğruluk (ZORUNLU):** Kaynak materyal yüklenmişse tüm somut bilgiler (tarih, kişi, olay) o kaynağa dayanmalıdır. Kaynak yoksa ya genel bilinen doğrulanmış olguları kullan ya da kurgusal karakterleri jenerik isimlerle ("dönemin bir sefaret kâtibi", "bölgeden geçen bir seyyah") ver; gerçek tarihî şahsiyetlere uydurma söz/rakam atfetme. Emin olmadığın tarih/rakam/alıntıyı icat etme, genel ama doğru bir ifade kullan.
"""

    secenek_blok = """
- **Seçenekler (ZORUNLU):** A, B, C, D, E olmak üzere 5 seçenek. Tüm seçenekler kelime sayısı, uzunluk ve dil yapısı bakımından birbirine YAKIN ve DENGELİ olmalı; doğru seçenek daha uzun/detaylı yazılıp görsel ipucu verilmemelidir. "Hepsi", "Hiçbiri", "A ve B" gibi kapsayıcı seçenekler KESİNLİKLE kullanılmamalı. Bağlam metnindeki cümleler seçeneklerde birebir tekrarlanmamalı (anlamca özdeş ama farklı kelimelerle). Çeldiriciler rastgele/bariz yanlış değil, konuyu eksik öğrenen öğrencinin düşeceği kavram yanılgılarından seçilmeli.
- Doğru cevap açıkça belirtilmeli ve kısa gerekçeli bir "Çözüm Açıklaması" eklenmelidir.
"""

    if baglam_temelli:
        kurgu_blok = f"""
### 🎯 BAĞLAM TEMELLİ SORU KURGUSU (TYMM Kılavuzu — ZORUNLU 5 ADIMLI SÜREÇ):

**ADIM 1 — Hedef:** Ölçülecek çıktı ve süreç bileşenleri yukarıda verilmiştir. Her soru bu süreç bileşenlerinden birini hedeflemelidir.

**ADIM 2 — Bağlamın Kurgulanması:**
- Bağlam metni TAM OLARAK İKİ (2) PARAGRAFTAN oluşmalıdır. İlk paragraf durumu/olayı/belgeyi tanıtır, ikinci paragraf ayrıntı/gelişme/farklı bakış açısı sunarak derinleştirir.
- Bir tarihçinin karşılaşacağı türden birincil/ikincil kaynak niteliğinde materyal (hatırat, dönemin gazetesi, arşiv belgesi, seyahatname, tarihçi değerlendirmesi, mektup vb.) kullanılmalı. Bağlam dekor olmamalı, gerçek bir çıkarım görevi sunmalı.
- Dar sosyoekonomik referanslardan kaçın; 11. sınıf düzeyine uygun, sade ve erişilebilir bir dil kullan.

**ADIM 3 — Sorular:**
- Bu TEK bağlam metnine dayanan TAM OLARAK {soru_sayisi} soru üretilmelidir; bu sayının altında kalınmamalıdır.
- Her soru farklı bir süreç bileşenini/bilişsel boyutu ölçmeli (bilgi/çıkarım, neden-sonuç, karşılaştırma, yargı, genelleme). Sorular birbirinin tekrarı olmamalı.
- **Soru Kökü Çeşitliliği (ZORUNLU):** Soruların soru kökü kalıpları da çeşitlenmeli; hepsi "hangi yargıya ulaşılabilir" veya "hangi konuya yönelmeli" gibi tek bir kalıpta olmamalıdır. Farklı köklerden yararlan: neden-sonuç ("...in temel nedeni aşağıdakilerden hangisidir?"), karşılaştırma ("...ile ...arasındaki temel fark nedir?"), çıkarım ("...den hangisi çıkarılabilir?"), yargı ("...değerlendirmelerinden hangisine ulaşılamaz?"). En az iki farklı kök kalıbı kullanılmalıdır.
- **İpucu Zinciri Yasağı:** Sorular arasında bağımlılık kurulmamalı; "bir önceki soruda bulduğunuz sonuca göre" tarzı ifadeler KULLANILMAMALI. Her soru bağlama bağımsız atıfta bulunmalı ("Bu parçaya göre...").
- Soru kökü: çift olumsuzluk ve "Sizce" gibi öznel ifadeler kullanılmamalı; "Metne göre" gibi nesnel ifadeler tercih edilmeli. Soru kökünde konu tekrar anlatılmamalı.
- Bağlam metni yalnızca bir kez en başta verilmeli; sorular altında 1., 2., 3. şeklinde sıralanmalı.
{secenek_blok}
**ADIM 4 — Güçlü Çeldiriciler:** Yanlış seçenekler metinle ilişkili ama ustaca yanıltıcı olmalı (kısmen doğru-eksik bilgi, ilgisiz ayrıntı, bağlam dışı doğru bilgi, mantık hatası).

**ADIM 5 — İşlevsellik Testi (ZORUNLU ÖZ-DENETİM):** Her soru için "öğrenci metni okumadan, sadece ön bilgisiyle veya seçenekleri eleyerek cevaplayabilir mi?" diye sor. Cevap "evet" ise soruyu, cevabın mutlaka metne dayanacağı şekilde yeniden kurgula.
"""
    else:
        kurgu_blok = f"""
### 🎯 NORMAL (KLASİK ÖSYM TARZI) SORU KURGUSU:
- TAM OLARAK {soru_sayisi} bağımsız soru üretilmelidir; her soru kendi kısa öncülüne/soru köküne sahip olabilir.
- Sorular yukarıdaki süreç bileşenlerini ölçmeli, kazanımın hedeflediği bilişsel beceriyi (analiz, çıkarım, karşılaştırma) doğrudan sınamalıdır.
- Uzun bağlam metni zorunlu değildir; ÖSYM tarzı doğrudan, net ve öz sorular hazırlanabilir. Yine de sorular ezber değil kavrayış ölçmelidir.
{secenek_blok}
- Çeldiriciler güçlü olmalı, konuyu eksik bilen öğrencinin düşeceği yanılgılardan seçilmelidir.
"""

    kaynak_blok = ""
    if kaynak_metin:
        kaynak_blok = f"""
### 📚 YÜKLENEN KAYNAK MATERYAL (ZORUNLU ÖNCELİK):
Bağlam ve soruları ÖNCELİKLE aşağıdaki kaynaklardaki bilgi, olay, kişi ve tarihlere dayanarak oluştur. Kaynakta kazanımla ilgili yeterli bilgi yoksa bunu belirtip genel tarihî bilgini kullan.

{kaynak_metin}
---
"""

    prompt = f"""Sen Türkiye Yüzyılı Maarif Modeli (TYMM) Bağlam Temelli Soru Yazım Kılavuzu'na ve ÖSYM standartlarına hakim, deneyimli bir Tarih öğretmeni/soru yazarısın.

Aşağıda TAM OLARAK belirtilen ünite, öğrenme çıktısı ve süreç bileşenleri için {soru_sayisi} adet {soru_kategorisi.lower()} soru üret. Bu parametrelerin DIŞINA çıkma; yalnızca verilen süreç bileşenlerini ölç.

---
### 📋 SORU PARAMETRELERİ
- **Ders:** Tarih (11. Sınıf)
- **Ünite:** {unite}
- **Öğrenme Çıktısı:** {cikti_tam}
- **Ölçülecek Süreç Bileşen(ler)i (alt başlıklar):**
{surec_listesi}
- **Soru Kategorisi:** {soru_kategorisi}
- **Zorluk Düzeyi:** {zorluk}
- **Üretilecek Soru Sayısı (KESİN):** {soru_sayisi}
{"- **Özel Bağlam Notu:** " + ek_baglam if ek_baglam else ""}
---
{program_blok}
{kaynak_blok}
### ✍️ YAZIM KURALLARI:
{uslup_blok}{dogruluk_blok}{kurgu_blok}
### ⏱️ ÇIKTI YÖNETİMİ:
İstenen soru sayısını tamamlamak, açıklamaları uzatmaktan HER ZAMAN önceliklidir. Yer daralırsa açıklamaları kısalt ama SAYIYI TAMAMLA; asla son soruyu yarım bırakma.

Çıktıyı okunaklı bir Markdown formatında, her soruyu numaralandırarak sun.
"""
    return prompt


# ==========================================
# OTURUM DURUMU
# ==========================================
if "uretilen_soru" not in st.session_state:
    st.session_state.uretilen_soru = None
if "son_uretim_meta" not in st.session_state:
    st.session_state.son_uretim_meta = None
if "kutuphane_yuklendi" not in st.session_state:
    metin, adlar = kutuphaneyi_diskten_yukle()
    st.session_state.kaynak_metin = metin
    st.session_state.kaynak_dosya_adlari = adlar
    st.session_state.kutuphane_yuklendi = True


# ==========================================
# BAŞLIK
# ==========================================
st.title("🏛️ TYMM 11. Sınıf Tarih Soru Üreteci")
st.markdown("**Türkiye Yüzyılı Maarif Modeli** müfredatına birebir bağlı, kademeli seçimli soru üretimi ve soru havuzu.")
if not DOCX_MEVCUT:
    st.warning("Word (.docx) indirme için `python-docx` gerekli: `pip install python-docx`")
st.divider()

# ==========================================
# SIDEBAR — KADEMELİ SEÇİM
# ==========================================
with st.sidebar:
    st.header("⚙️ Müfredat Seçimi")
    st.caption("Sırasıyla ünite → öğrenme çıktısı → süreç bileşeni (alt başlık) seçin.")

    unite_secimi = st.selectbox("1️⃣ Ünite:", options=list(MUFREDAT.keys()))

    ciktilar = MUFREDAT[unite_secimi]
    cikti_kod = st.selectbox(
        "2️⃣ Öğrenme Çıktısı:",
        options=list(ciktilar.keys()),
        format_func=lambda k: f"{k} — {ciktilar[k]['tam'].split('.', 3)[-1].strip()[:38]}..."
    )
    cikti_tam = ciktilar[cikti_kod]["tam"]

    surecler_dict = ciktilar[cikti_kod]["surecler"]
    surec_secimi = st.multiselect(
        "3️⃣ Süreç Bileşen(ler)i (alt başlık):",
        options=list(surecler_dict.keys()),
        default=list(surecler_dict.keys())[:1],
        format_func=lambda s: s[:60] + ("..." if len(s) > 60 else "")
    )

    st.divider()
    st.header("🎯 Soru Formatı")

    soru_kategorisi = st.radio("Soru Kategorisi:", ["Bağlam Temelli", "Normal"], horizontal=True)
    zorluk = st.radio("Zorluk Düzeyi:", ZORLUK_SECENEKLERI, horizontal=True)

    onerilen = 0
    anahtar = "bt" if soru_kategorisi == "Bağlam Temelli" else "normal"
    for s in surec_secimi:
        onerilen += surecler_dict[s][anahtar].get(zorluk, 0)

    if onerilen > 0:
        st.info(f"📊 Müfredat tablosuna göre önerilen soru sayısı: **{onerilen}**")

    varsayilan_sayi = max(onerilen, 1) if soru_kategorisi == "Normal" else max(onerilen, 5)
    soru_sayisi = st.number_input(
        "Üretilecek Soru Sayısı:",
        min_value=1, max_value=30, value=int(varsayilan_sayi), step=1,
        help="Varsayılan değer müfredat tablosundaki öneriden gelir; değiştirebilirsiniz."
    )
    if soru_kategorisi == "Bağlam Temelli" and soru_sayisi < 5:
        st.caption("ℹ️ Bağlam temelli sorularda aynı metne dayalı en az 5 soru önerilir.")
    if soru_sayisi > 10:
        st.caption("⚠️ 10+ soruda model token sınırına takılabilir; büyük setleri gruplar hâlinde ürettirin.")

    st.divider()
    st.header("🤖 AI ile Üretim")
    saglayici = st.radio("Sağlayıcı:",
                         ["Anthropic (Claude)", "Google (Gemini)", "DeepSeek"],
                         horizontal=True)

    if saglayici == "Anthropic (Claude)":
        if not ANTHROPIC_MEVCUT:
            st.warning("`anthropic` kurulu değil: `pip install anthropic`")
        api_key = st.text_input("Anthropic API Anahtarı:", type="password")
        model_secimi = st.selectbox("Model:", CLAUDE_MODELLERI)
    elif saglayici == "Google (Gemini)":
        if not GEMINI_MEVCUT:
            st.warning("`google-generativeai` kurulu değil: `pip install google-generativeai`")
        api_key = st.text_input("Gemini API Anahtarı:", type="password")
        model_secimi = st.selectbox("Model:", GEMINI_MODELLERI)
    else:  # DeepSeek
        if not OPENAI_MEVCUT:
            st.warning("`openai` kurulu değil: `pip install openai`")
        api_key = st.text_input("DeepSeek API Anahtarı:", type="password",
                                help="platform.deepseek.com üzerinden alınır. Katı dakikalık istek limiti yoktur ve çok ucuzdur.")
        model_secimi = st.selectbox("Model:", DEEPSEEK_MODELLERI)
        st.caption("💡 DeepSeek'te dakikalık istek kotası yoktur; büyük taramalar için uygundur.")


# ==========================================
# ANA EKRAN — SEÇİM ÖZETİ
# ==========================================
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("📌 Seçili Müfredat Öğeleri")
    st.markdown(f"**Ünite:** {unite_secimi}")
    st.markdown(f"**Öğrenme Çıktısı:** `{cikti_kod}`")
    st.caption(cikti_tam)
    if surec_secimi:
        st.markdown("**Seçili Süreç Bileşenleri:**")
        for s in surec_secimi:
            st.markdown(f"- {s}")
    else:
        st.warning("Lütfen en az bir süreç bileşeni (alt başlık) seçin.")

    st.subheader("📝 Özel Bağlam Notu (İsteğe Bağlı)")
    ek_baglam = st.text_area(
        "Soruda geçmesini istediğiniz özel tarihçi, belge veya olay:",
        placeholder="Örn: Karlofça müzakereleri bağlamı kullanılsın...",
        height=80
    )

with col2:
    st.subheader("📊 Bu Çıktının Soru Dağılımı")
    st.caption("Müfredat tablosu (Excel) — seçili çıktının tüm süreç bileşenleri toplamı:")
    top_bt = {z: 0 for z in ZORLUK_SECENEKLERI}
    top_nrm = {z: 0 for z in ZORLUK_SECENEKLERI}
    for s, v in surecler_dict.items():
        for z in ZORLUK_SECENEKLERI:
            top_bt[z] += v["bt"].get(z, 0)
            top_nrm[z] += v["normal"].get(z, 0)
    st.markdown("**Bağlam Temelli**")
    st.markdown(f"Kolay: {top_bt['Kolay']} · Orta: {top_bt['Orta']} · Zor: {top_bt['Zor']}")
    st.markdown("**Normal**")
    st.markdown(f"Kolay: {top_nrm['Kolay']} · Orta: {top_nrm['Orta']} · Zor: {top_nrm['Zor']}")

st.divider()

# ==========================================
# ANA EKRAN — KAYNAK KÜTÜPHANESİ
# ==========================================
st.subheader("📚 Kaynak Kitap / Doküman Kütüphanesi")
st.caption("İstediğiniz kadar PDF veya .txt yükleyebilirsiniz. Kaynaklar kalıcı olarak diske kaydedilir ve soru üretiminde birincil kaynak olur.")

yeni_dosyalar = st.file_uploader(
    "PDF / metin dosyalarını sürükleyip bırakın:",
    type=["pdf", "txt"], accept_multiple_files=True, key="kutuphane_yukleyici"
)
st.caption(f"📁 Kayıt yeri: `{KUTUPHANE_KLASORU}`")

maks_kaynak_karakter = st.slider(
    "Modele gönderilecek azami kaynak metni (karakter):",
    min_value=20000, max_value=800000, value=MAKS_KAYNAK_KARAKTER_VARSAYILAN, step=10000,
    help="1 token ≈ 4 karakter. Yüksek değerler API maliyetini/süreyi artırır ve "
         "ücretsiz katmanda dakikalık token kotasını (429 hatası) hızla doldurur."
)
if maks_kaynak_karakter > 200000:
    st.caption("⚠️ Yüksek kaynak metni ücretsiz katmanda '429 kota aşımı' hatasına yol açabilir. "
               "Sık hata alıyorsanız bu değeri 40.000–60.000 aralığına düşürün.")

cy, ct = st.columns(2)
with cy:
    if st.button("➕ Kütüphaneye Kaydet", use_container_width=True, disabled=not yeni_dosyalar):
        with st.spinner("Kaydediliyor..."):
            yuklenen_dosyalari_isle(yeni_dosyalar)
            st.session_state.kaynak_metin, st.session_state.kaynak_dosya_adlari = kutuphaneyi_diskten_yukle()
        st.success(f"{len(yeni_dosyalar)} dosya kaydedildi.")
        st.rerun()
with ct:
    if st.button("🗑️ Kütüphaneyi Temizle", use_container_width=True, disabled=not st.session_state.kaynak_dosya_adlari):
        for d in KUTUPHANE_KLASORU.glob("*"):
            d.unlink(missing_ok=True)
        st.session_state.kaynak_metin, st.session_state.kaynak_dosya_adlari = "", []
        st.success("Temizlendi.")
        st.rerun()

if st.session_state.kaynak_dosya_adlari:
    tk = len(st.session_state.kaynak_metin)
    st.write(f"**Kütüphane ({len(st.session_state.kaynak_dosya_adlari)} dosya, ~{tk:,} karakter):**")
    for i, ad in enumerate(st.session_state.kaynak_dosya_adlari):
        c1, c2 = st.columns([5, 1])
        c1.write(f"📄 {ad}")
        if c2.button("Kaldır", key=f"kaldir_{i}"):
            (KUTUPHANE_KLASORU / ad).unlink(missing_ok=True)
            st.session_state.kaynak_metin, st.session_state.kaynak_dosya_adlari = kutuphaneyi_diskten_yukle()
            st.rerun()
    if tk > maks_kaynak_karakter:
        st.warning(f"Kütüphane {tk:,} karakter; yalnızca ilk {maks_kaynak_karakter:,} karakter modele gönderilecek. Sınırı yukarıdan artırabilirsiniz.")
else:
    st.caption("Henüz kaynak eklenmedi. Kaynak yoksa sorular genel tarih bilgisiyle üretilir.")

st.divider()

# ==========================================
# ANA EKRAN — PROMPT VE ÜRETİM
# ==========================================
st.subheader("🚀 Soru Üretimi")

if surec_secimi:
    generated_prompt = build_prompt(
        unite_secimi, cikti_kod, cikti_tam, surec_secimi, soru_kategorisi,
        zorluk, soru_sayisi,
        kaynak_metin=st.session_state.kaynak_metin[:maks_kaynak_karakter],
        ek_baglam=ek_baglam
    )
else:
    generated_prompt = "⚠️ Lütfen sol menüden en az bir süreç bileşeni (alt başlık) seçin."

with st.expander("🔍 Oluşturulan promptu görüntüle / kopyala"):
    st.text_area("Prompt:", value=generated_prompt, height=300)

b0, b1, b2 = st.columns(3)
with b0:
    uret_tiklandi = st.button(
        "✨ Soruyu Şimdi Üret", use_container_width=True, type="primary",
        disabled=not (surec_secimi and (ANTHROPIC_MEVCUT or GEMINI_MEVCUT or OPENAI_MEVCUT))
    )
with b1:
    st.download_button("⬇️ Promptu İndir (.txt)", data=generated_prompt,
                       file_name=f"prompt_{cikti_kod}.txt", mime="text/plain", use_container_width=True)
with b2:
    st.download_button("⬇️ Yapılandırma (.json)",
                       data=json.dumps({
                           "unite": unite_secimi, "cikti_kod": cikti_kod, "cikti_tam": cikti_tam,
                           "surecler": surec_secimi, "soru_kategorisi": soru_kategorisi,
                           "zorluk": zorluk, "soru_sayisi": soru_sayisi
                       }, ensure_ascii=False, indent=2),
                       file_name=f"config_{cikti_kod}.json", mime="application/json", use_container_width=True)

if uret_tiklandi:
    if saglayici == "Anthropic (Claude)" and not ANTHROPIC_MEVCUT:
        st.error("`anthropic` kurulu değil.")
    elif saglayici == "Google (Gemini)" and not GEMINI_MEVCUT:
        st.error("`google-generativeai` kurulu değil.")
    elif saglayici == "DeepSeek" and not OPENAI_MEVCUT:
        st.error("`openai` kurulu değil. `pip install openai` çalıştırın.")
    elif not api_key:
        st.error(f"Lütfen {saglayici} API anahtarınızı girin.")
    else:
        durum = st.status("Soru üretiliyor...", expanded=False)

        def ilerleme_bildir(msg):
            durum.update(label=msg, state="running")

        try:
            if saglayici == "Anthropic (Claude)":
                sonuc = soru_uret_api(api_key, model_secimi, generated_prompt, ilerleme=ilerleme_bildir)
            elif saglayici == "Google (Gemini)":
                sonuc = soru_uret_gemini(api_key, model_secimi, generated_prompt, ilerleme=ilerleme_bildir)
            else:  # DeepSeek
                sonuc = soru_uret_deepseek(api_key, model_secimi, generated_prompt, ilerleme=ilerleme_bildir)
            durum.update(label="Üretim tamamlandı.", state="complete")
            st.session_state.uretilen_soru = sonuc
            st.session_state.son_uretim_meta = {
                "unite": unite_secimi, "cikti_kod": cikti_kod, "cikti_tam": cikti_tam,
                "surecler": surec_secimi, "soru_kategorisi": soru_kategorisi,
                "zorluk": zorluk, "soru_sayisi": soru_sayisi,
                "model": model_secimi, "zaman": datetime.now().strftime("%Y-%m-%d %H:%M")
            }
        except Exception as e:
            durum.update(label="Üretim başarısız.", state="error")
            mesaj = str(e)
            if ANTHROPIC_MEVCUT and isinstance(e, anthropic.AuthenticationError):
                st.error("API anahtarı geçersiz görünüyor.")
            elif "429" in mesaj or "quota" in mesaj.lower() or "rate" in mesaj.lower():
                st.error("⏳ Kota/hız limiti aşıldı ve otomatik denemeler de yetmedi. "
                         "Birkaç dakika bekleyip tekrar deneyin. İpucu: kaynak metni sınırını "
                         "düşürmek (yandaki kaydırıcı) veya soru sayısını azaltmak limiti korur.")
            else:
                st.error(f"Hata oluştu: {e}")

if st.session_state.uretilen_soru:
    st.divider()
    st.subheader("📄 Üretilen Soru(lar)")
    st.markdown(st.session_state.uretilen_soru)

    k1, k2, k3 = st.columns(3)
    with k1:
        if st.button("💾 Soru Havuzuna Kaydet", use_container_width=True, type="primary"):
            meta = st.session_state.son_uretim_meta or {}
            kayit = {
                "id": datetime.now().strftime("%Y%m%d%H%M%S%f"),
                "unite": meta.get("unite", unite_secimi),
                "cikti_kod": meta.get("cikti_kod", cikti_kod),
                "cikti_tam": meta.get("cikti_tam", cikti_tam),
                "surecler": meta.get("surecler", surec_secimi),
                "soru_kategorisi": meta.get("soru_kategorisi", soru_kategorisi),
                "zorluk": meta.get("zorluk", zorluk),
                "soru_sayisi": meta.get("soru_sayisi", soru_sayisi),
                "model": meta.get("model", model_secimi),
                "zaman": meta.get("zaman", datetime.now().strftime("%Y-%m-%d %H:%M")),
                "icerik": st.session_state.uretilen_soru
            }
            havuza_kaydet(kayit)
            st.success("Soru havuzuna kaydedildi! Aşağıdaki havuzdan erişebilirsiniz.")
    with k2:
        st.download_button("⬇️ İndir (.md)", data=st.session_state.uretilen_soru,
                           file_name=f"sorular_{cikti_kod}_{zorluk}.md", mime="text/markdown", use_container_width=True)
    with k3:
        if DOCX_MEVCUT:
            meta = st.session_state.son_uretim_meta or {}
            ust = " · ".join(str(x) for x in [meta.get("cikti_kod", ""), meta.get("soru_kategorisi", ""),
                                              meta.get("zorluk", ""), meta.get("zaman", "")] if str(x).strip())
            word_tampon = ham_metin_word(st.session_state.uretilen_soru, ust_bilgi=ust)
            st.download_button(
                "⬇️ Word İndir (.docx)", data=word_tampon,
                file_name=f"sorular_{cikti_kod}_{zorluk}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.button("⬇️ Word (.docx)", disabled=True, use_container_width=True,
                      help="python-docx kurulu değil")

st.divider()

# ==========================================
# ANA EKRAN — SORU HAVUZU
# ==========================================
st.subheader("🗂️ Soru Havuzu")
st.caption("Kaydedilen tüm sorular ünite / öğrenme çıktısı / alt başlık bilgisiyle burada saklanır ve istendiğinde erişilir.")

havuz = havuzu_yukle()

if not havuz:
    st.info("Henüz havuza kaydedilmiş soru yok. Yukarıda soru üretip 'Soru Havuzuna Kaydet' butonuna basın.")
else:
    f1, f2, f3 = st.columns(3)
    with f1:
        f_unite = st.selectbox("Üniteye göre:", ["(Tümü)"] + list(MUFREDAT.keys()), key="f_unite")
    with f2:
        cikti_secenek = ["(Tümü)"] + (list(MUFREDAT[f_unite].keys()) if f_unite != "(Tümü)" else [])
        f_cikti = st.selectbox("Çıktıya göre:", cikti_secenek, key="f_cikti")
    with f3:
        f_kat = st.selectbox("Kategoriye göre:", ["(Tümü)", "Bağlam Temelli", "Normal"], key="f_kat")

    filtreli = havuz
    if f_unite != "(Tümü)":
        filtreli = [k for k in filtreli if k.get("unite") == f_unite]
    if f_cikti != "(Tümü)":
        filtreli = [k for k in filtreli if k.get("cikti_kod") == f_cikti]
    if f_kat != "(Tümü)":
        filtreli = [k for k in filtreli if k.get("soru_kategorisi") == f_kat]

    st.write(f"**{len(filtreli)} / {len(havuz)} kayıt gösteriliyor.**")

    # --- Toplu dışa aktarma butonları ---
    d_json, d_word_filt, d_word_tum = st.columns(3)
    with d_json:
        st.download_button(
            "⬇️ Tüm Havuz (.json)",
            data=json.dumps(havuz, ensure_ascii=False, indent=2),
            file_name="soru_havuzu.json", mime="application/json", use_container_width=True
        )
    with d_word_filt:
        if DOCX_MEVCUT and filtreli:
            st.download_button(
                f"⬇️ Filtreli → Word ({len(filtreli)})",
                data=coklu_kayit_word(filtreli, baslik="TYMM Tarih — Seçili Sorular"),
                file_name="sorular_filtreli.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.button("⬇️ Filtreli → Word", disabled=True, use_container_width=True,
                      help="python-docx kurulu değil veya kayıt yok")
    with d_word_tum:
        if DOCX_MEVCUT and havuz:
            st.download_button(
                f"⬇️ Tüm Havuz → Word ({len(havuz)})",
                data=coklu_kayit_word(havuz, baslik="TYMM Tarih — Tüm Soru Havuzu"),
                file_name="soru_havuzu_tamami.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.button("⬇️ Tüm Havuz → Word", disabled=True, use_container_width=True,
                      help="python-docx kurulu değil")

    for kayit in filtreli:
        baslik = f"📘 {kayit.get('cikti_kod','')} · {kayit.get('soru_kategorisi','')} · {kayit.get('zorluk','')} · {kayit.get('soru_sayisi','')} soru · {kayit.get('zaman','')}"
        with st.expander(baslik):
            st.caption(f"**Ünite:** {kayit.get('unite','')}")
            st.caption(f"**Çıktı:** {kayit.get('cikti_tam','')}")
            if kayit.get("surecler"):
                st.caption("**Alt başlıklar:** " + "; ".join(kayit["surecler"]))
            st.markdown("---")
            st.markdown(kayit.get("icerik", ""))
            d1, d2, d3 = st.columns([1, 1, 4])
            with d1:
                st.download_button("⬇️ .md", data=kayit.get("icerik", ""),
                                   file_name=f"soru_{kayit.get('cikti_kod','')}_{kayit.get('id','')}.md",
                                   mime="text/markdown", key=f"dl_{kayit['id']}")
            with d2:
                if DOCX_MEVCUT:
                    st.download_button(
                        "⬇️ .docx",
                        data=tek_kayit_word(kayit),
                        file_name=f"soru_{kayit.get('cikti_kod','')}_{kayit.get('id','')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dlw_{kayit['id']}"
                    )
                else:
                    st.button("⬇️ .docx", disabled=True, key=f"dlw_{kayit['id']}", help="python-docx yok")
            with d3:
                if st.button("🗑️ Bu kaydı havuzdan sil", key=f"sil_{kayit['id']}"):
                    havuzdan_sil(kayit["id"])
                    st.rerun()
